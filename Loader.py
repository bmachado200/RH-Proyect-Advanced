import os
import re
import logging
import warnings
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

import chromadb
from chromadb.utils.embedding_functions import OpenAIEmbeddingFunction
from dotenv import load_dotenv # Used for loading API key from .env in main execution
import tiktoken

# Document processing libraries
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text
# import PyPDF2 # Fallback for PDF, pdfminer.six is generally preferred

# OCR specific imports
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image # Pillow for image manipulation
    OCR_CAPABLE = True
except ImportError:
    OCR_CAPABLE = False
    print("WARNING: OCR libraries (pytesseract, pdf2image, Pillow) not installed. PDF OCR functionality will be disabled.")
    print("To enable OCR, please install them: pip install pytesseract pdf2image Pillow")
    print("And ensure Tesseract OCR engine is installed on your system.")

# ==============================================================================
# SECCIÓN 2: CONFIGURACIÓN
# ==============================================================================
# Define todas las constantes y parámetros configurables para el script.
# Esto incluye rutas de directorios, el nombre de la colección en la base de datos,
# la clave de API de OpenAI, y parámetros para la división del texto en fragmentos (chunks).

load_dotenv() # Load environment variables from .env file

BASE_DIR = Path(__file__).resolve().parent
HR_DOCS_DIR = BASE_DIR / "HR"  # Directory where your .pdf and .docx files are
CHROMA_DB_DIR = BASE_DIR / "chroma_db_hr_ocr" # Persistent storage for ChromaDB

COLLECTION_NAME = "hr_documents_ocr_production_v1" # Ensure this matches query.py

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EMBEDDING_MODEL_NAME = "text-embedding-3-large" # Ensure this matches query.py

# Chunking Parameters
TARGET_CHUNK_CHAR_SIZE = 1500  # Target character size for chunks
CHUNK_CHAR_OVERLAP = 300     # Number of characters to overlap between chunks
MAX_TOKENS_PER_CHUNK = 8000  # Maximum tokens per chunk (hard limit for embedding model)

# Tesseract OCR Configuration (Update path if Tesseract is not in your system PATH)
# Example for Windows:
# TESSERACT_CMD_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# Example for Linux/macOS (if installed via package manager, usually not needed):
TESSERACT_CMD_PATH = None # Set this if pytesseract cannot find tesseract

if OCR_CAPABLE and TESSERACT_CMD_PATH:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD_PATH

# Poppler path for pdf2image (Windows specific, often needed if Poppler is not in PATH)
# POPPLER_PATH = r"C:\path\to\poppler-xxx\bin" # Example for Windows
POPPLER_PATH = None


# ==============================================================================
# SECCIÓN 3: CONFIGURACIÓN DE LOGGING
# ==============================================================================
# Configura el sistema de registro (logging) para mostrar mensajes informativos,
# advertencias y errores durante la ejecución del script, facilitando el seguimiento
# y la depuración.

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(module)s - %(message)s')
logger = logging.getLogger(__name__)
warnings.filterwarnings("ignore", category=UserWarning, module='pdfminer')
pil_logger = logging.getLogger('PIL') # Suppress verbose PIL logging if necessary
pil_logger.setLevel(logging.INFO)

# ==============================================================================
# SECCIÓN 4: FUNCIONES DE AYUDA Y TOKENIZADOR
# ==============================================================================
# Contiene funciones de utilidad, como la que inicializa la función de embedding
# de OpenAI y la que cuenta tokens en un texto usando 'tiktoken'.

def get_embedding_function() -> OpenAIEmbeddingFunction:
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API_KEY not found in environment variables.")
    return OpenAIEmbeddingFunction(api_key=OPENAI_API_KEY, model_name=EMBEDDING_MODEL_NAME)

try:
    tokenizer = tiktoken.encoding_for_model(EMBEDDING_MODEL_NAME)
except KeyError:
    logger.warning(f"Encoding for {EMBEDDING_MODEL_NAME} not found. Using cl100k_base.")
    tokenizer = tiktoken.get_encoding("cl100k_base")

def count_tokens(text: str) -> int:
    return len(tokenizer.encode(text))

# ==============================================================================
# SECCIÓN 5: EXTRACCIÓN DE TEXTO CON OCR
# ==============================================================================
# Funciones dedicadas a extraer texto de diferentes tipos de archivo.
# 'extract_text_from_pdf_with_ocr' es notable por su estrategia dual: primero
# intenta extraer texto digital y, si falla o el texto es pobre, recurre al
# Reconocimiento Óptico de Caracteres (OCR).

def ocr_pdf_page(image: Image) -> str:
    """Performs OCR on a single image (PDF page)."""
    if not OCR_CAPABLE:
        logger.warning("OCR called but not capable. Returning empty string.")
        return ""
    try:
        # Specify language(s) for Tesseract. 'spa' for Spanish, 'eng' for English.
        custom_oem_psm_config = r'--oem 3 --psm 6 -l spa+eng' # Adjust as needed
        text = pytesseract.image_to_string(image, config=custom_oem_psm_config)
        return text
    except Exception as e:
        logger.error(f"Tesseract OCR error on page: {e}")
        return ""

def extract_text_from_pdf_with_ocr(file_path: Path, force_ocr: bool = False) -> str:
    logger.info(f"Extracting text from PDF (Force OCR: {force_ocr}): {file_path.name}")
    digital_text = ""
    
    if not force_ocr:
        try:
            digital_text = pdfminer_extract_text(str(file_path))
            # Basic check for meaningful content
            if digital_text and len(digital_text.strip()) > 200:
                logger.info(f"Successfully extracted digital text from {file_path.name} using pdfminer.six. Length: {len(digital_text)} chars.")
                return digital_text.strip()
            else:
                logger.info(f"Digital text extraction from {file_path.name} was minimal or empty. Proceeding with OCR attempt.")
        except Exception as e:
            logger.warning(f"Error with digital text extraction for {file_path.name} (pdfminer): {e}. Proceeding with OCR.")

    if not OCR_CAPABLE:
        logger.warning(f"OCR is not available. Returning any digitally extracted text for {file_path.name}.")
        return digital_text.strip()

    logger.info(f"Performing OCR on {file_path.name}...")
    ocr_text_parts = []
    try:
        if POPPLER_PATH and os.name == 'nt': # Windows check for Poppler path
            images = convert_from_path(str(file_path), poppler_path=POPPLER_PATH)
        else:
            images = convert_from_path(str(file_path)) # Assumes Poppler is in PATH for other OS or if POPPLER_PATH is None
        
        if not images:
            logger.warning(f"pdf2image could not convert {file_path.name} to images. No OCR possible.")
            return digital_text.strip() # Fallback

        for i, page_image in enumerate(images):
            logger.debug(f"OCR'ing page {i+1} of {len(images)} for {file_path.name}...")
            page_ocr_text = ocr_pdf_page(page_image)
            if page_ocr_text:
                ocr_text_parts.append(page_ocr_text)
        
        full_ocr_text = "\n\n[End of Page]\n\n".join(ocr_text_parts) # Add page breaks
        
        if full_ocr_text.strip():
            logger.info(f"Successfully extracted text from {file_path.name} using OCR. Length: {len(full_ocr_text)} chars.")
            return full_ocr_text.strip()
        else:
            logger.warning(f"OCR yielded no text for {file_path.name}. Falling back to any digital text found.")
            return digital_text.strip()

    except Exception as e:
        logger.error(f"Error during OCR processing for {file_path.name}: {e}")
        return digital_text.strip() # Fallback

def extract_text_from_docx(file_path: Path) -> str:
    logger.info(f"Extracting text from DOCX: {file_path.name}")
    try:
        doc = Document(file_path)
        full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        extracted_content = "\n\n".join(full_text) # Use double newline for paragraphs
        logger.info(f"Extracted text from DOCX {file_path.name}. Length: {len(extracted_content)} chars.")
        return extracted_content
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path.name}: {e}")
        return ""

# ==============================================================================
# SECCIÓN 6: PROCESAMIENTO DE TEXTO, CHUNKING Y METADATOS
# ==============================================================================
# Este bloque se encarga de procesar el texto extraído: limpiarlo, dividirlo en
# fragmentos (chunks) de tamaño manejable para el modelo de embedding, y generar
# metadatos detallados para cada archivo y chunk.

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text) # Consolidate multiple whitespace characters
    text = re.sub(r'\n\s*\n', '\n\n', text) # Normalize multiple newlines to consistent double newlines
    return text.strip()

def split_text_into_chunks(text: str, file_name: str) -> List[str]:
    logger.debug(f"Splitting text for {file_name}, original length: {len(text)} chars, ~{count_tokens(text)} tokens")
    
    # Split by double newlines, assuming these are paragraph breaks after cleaning
    paragraphs = text.split('\n\n')
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks: List[str] = []
    current_chunk_texts: List[str] = []
    current_char_count = 0
    
    for i, para in enumerate(paragraphs):
        para_char_count = len(para)
        # Check if adding the current paragraph would exceed size or token limits
        # +2 for potential "\n\n" joiner
        would_exceed_char_size = current_char_count + para_char_count + (len(current_chunk_texts) * 2) > TARGET_CHUNK_CHAR_SIZE
        
        # Temporarily form candidate chunk to check token count
        candidate_chunk_text = "\n\n".join(current_chunk_texts + [para]) if current_chunk_texts else para
        candidate_tokens = count_tokens(candidate_chunk_text)
        would_exceed_token_limit = candidate_tokens > MAX_TOKENS_PER_CHUNK

        if current_chunk_texts and (would_exceed_char_size or would_exceed_token_limit):
            # Finalize the current chunk
            final_chunk_text = "\n\n".join(current_chunk_texts)
            
            if count_tokens(final_chunk_text) > MAX_TOKENS_PER_CHUNK:
                logger.warning(f"Chunk for {file_name} (char size {len(final_chunk_text)}) was over token limit BEFORE adding new para. Splitting by tokens.")
                token_sub_chunks = split_chunk_by_tokens_recursive(final_chunk_text, MAX_TOKENS_PER_CHUNK)
                chunks.extend(token_sub_chunks)
            else:
                chunks.append(final_chunk_text)
            
            # Start new chunk, considering overlap
            if CHUNK_CHAR_OVERLAP > 0 and final_chunk_text:
                # Get the last few paragraphs or characters for overlap
                # This is a simplified overlap; more sophisticated methods exist.
                # Here, we take the last paragraph of the previous chunk if it's substantial,
                # or a character slice.
                overlap_text = final_chunk_text[-CHUNK_CHAR_OVERLAP:]
                # Try to find a sensible break for overlap text
                last_para_break = overlap_text.rfind("\n\n")
                if last_para_break != -1 and len(overlap_text[last_para_break:].strip()) > 50 : # Meaningful paragraph
                     current_chunk_texts = [overlap_text[last_para_break:].strip(), para]
                else: # Fallback to char overlap or just the new para
                     current_chunk_texts = [para] # Simplest: start new chunk with current para, can be improved.
            else:
                current_chunk_texts = [para]
            current_char_count = sum(len(p) + 2 for p in current_chunk_texts) -2 if current_chunk_texts else 0
        
        else: # Add paragraph to current chunk if it doesn't exceed limits on its own
            if would_exceed_token_limit and not current_chunk_texts: # Current paragraph itself is too large
                logger.warning(f"Single paragraph in {file_name} (char size {para_char_count}) exceeds token limit. Splitting by tokens.")
                token_sub_chunks = split_chunk_by_tokens_recursive(para, MAX_TOKENS_PER_CHUNK)
                chunks.extend(token_sub_chunks)
                current_chunk_texts = [] # Reset as this paragraph has been processed
                current_char_count = 0
            else:
                current_chunk_texts.append(para)
                current_char_count += para_char_count + 2 # +2 for potential "\n\n"

    # Add the last remaining chunk
    if current_chunk_texts:
        final_chunk_text = "\n\n".join(current_chunk_texts)
        if count_tokens(final_chunk_text) > MAX_TOKENS_PER_CHUNK:
            logger.warning(f"Final chunk for {file_name} (char size {len(final_chunk_text)}) exceeded token limit. Splitting by tokens.")
            token_sub_chunks = split_chunk_by_tokens_recursive(final_chunk_text, MAX_TOKENS_PER_CHUNK)
            chunks.extend(token_sub_chunks)
        else:
            chunks.append(final_chunk_text)
            
    logger.info(f"Split {file_name} into {len(chunks)} chunks.")
    if not chunks and text: # If text existed but no chunks were made (e.g. single huge paragraph)
        logger.warning(f"Original text for {file_name} was not empty but resulted in 0 primary chunks. This might happen if a single paragraph was split by tokens.")

    return chunks

def split_chunk_by_tokens_recursive(text: str, max_tokens: int) -> List[str]:
    """Splits text into sub-chunks if it exceeds max_tokens, recursively."""
    if count_tokens(text) <= max_tokens:
        return [text]

    logger.debug(f"Recursively splitting chunk by tokens (original tokens: {count_tokens(text)})")
    # Simple split: find midpoint by characters (approximation) and then decode.
    # More advanced: split at sentence boundaries using the tokenizer.
    tokens = tokenizer.encode(text)
    sub_chunks: List[str] = []
    
    current_pos = 0
    while current_pos < len(tokens):
        end_pos = min(current_pos + max_tokens, len(tokens))
        # Try to find a sentence break near end_pos if not at the very end
        # This is a simplified approach; proper sentence tokenization is more robust
        if end_pos < len(tokens):
            temp_decode_for_break = tokenizer.decode(tokens[current_pos:end_pos])
            last_sentence_end = -1
            for marker in ['. ', '! ', '? ', '\n\n']: # Prefer paragraph breaks
                idx = temp_decode_for_break.rfind(marker)
                if idx > last_sentence_end:
                    last_sentence_end = idx + len(marker) # Include the marker
            
            if last_sentence_end > len(temp_decode_for_break) / 2 : # If break is reasonably far in
                 # Re-encode up to the break to get accurate token count for this segment
                adjusted_segment_tokens = tokenizer.encode(temp_decode_for_break[:last_sentence_end])
                end_pos = current_pos + len(adjusted_segment_tokens)


        chunk_tokens = tokens[current_pos:end_pos]
        decoded_chunk = tokenizer.decode(chunk_tokens)
        if decoded_chunk.strip(): # Avoid empty chunks
            sub_chunks.append(decoded_chunk)
        current_pos = end_pos
        
    logger.debug(f"Resulted in {len(sub_chunks)} token-based sub-chunks.")
    return sub_chunks


def get_file_metadata(file_path: Path) -> Dict[str, Any]:
    stat = file_path.stat()
    return {
        "file_name": file_path.name,
        "file_path_str": str(file_path), # Using a different key to avoid conflict if 'file_path' is used by ChromaDB
        "absolute_path_str": str(file_path.resolve()),
        "file_type": file_path.suffix.lower(),
        "file_size_bytes": stat.st_size,
        "created_at_timestamp": stat.st_ctime,
        "modified_at_timestamp": stat.st_mtime, # Timestamp of last modification
        "processed_at_loader_timestamp": datetime.now().timestamp()
    }

# ==============================================================================
# SECCIÓN 7: INTERACCIÓN CON CHROMADB
# ==============================================================================
# Funciones para inicializar el cliente de ChromaDB, crear o acceder a una
# colección, y la función principal 'process_and_load_documents' que orquesta
# todo el proceso de carga de datos en la base de datos vectorial.

def initialize_chroma_client() -> chromadb.ClientAPI: # Or simply chromadb.Client    logger.info(f"Initializing ChromaDB client at: {CHROMA_DB_DIR}")
    if not CHROMA_DB_DIR.exists():
        CHROMA_DB_DIR.mkdir(parents=True, exist_ok=True)
    return chromadb.PersistentClient(path=str(CHROMA_DB_DIR))

def get_or_create_collection(client: chromadb.ClientAPI, embedding_fx: OpenAIEmbeddingFunction) -> chromadb.Collection:
    logger.info(f"Getting or creating ChromaDB collection: {COLLECTION_NAME}")
    return client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=embedding_fx, # Critical for both adding and querying
        metadata={"hnsw:space": "cosine"} # Using cosine distance
    )

def process_and_load_documents(collection_to_load: chromadb.Collection, force_ocr_all_pdfs: bool = False, force_reprocess_all_files: bool = False):
    logger.info(f"Starting document processing from: {HR_DOCS_DIR}. Force OCR all PDFs: {force_ocr_all_pdfs}. Force reprocess all: {force_reprocess_all_files}")
    if not HR_DOCS_DIR.exists():
        logger.error(f"Documents directory not found: {HR_DOCS_DIR}. Please create it and add .pdf or .docx files.")
        HR_DOCS_DIR.mkdir(parents=True, exist_ok=True) # Create if not exists
        logger.info(f"{HR_DOCS_DIR} created. Please add documents and re-run.")
        return

    processed_files_count = 0
    new_chunks_added_this_run = 0

    for item in HR_DOCS_DIR.iterdir():
        if item.is_file():
            file_path = item
            file_name = item.name
            file_suffix = item.suffix.lower()
            logger.info(f"--- Processing file: {file_name} ---")

            if not (file_suffix == ".pdf" or file_suffix == ".docx"):
                logger.info(f"Skipping unsupported file type: {file_name}")
                continue

            current_file_metadata = get_file_metadata(file_path)
            
            # Check if file needs update
            needs_update = True # Assume it needs update by default
            if not force_reprocess_all_files:
                # Query for chunks from this file_name and check their 'modified_at_timestamp'
                # This assumes 'file_name' is unique and 'modified_at_timestamp' is stored reliably.
                try:
                    existing_data = collection_to_load.get(
                        where={"file_name": file_name}, # Query by 'file_name' stored in metadata
                        include=["metadatas"]
                    )
                    if existing_data and existing_data['ids']:
                        # Check the mod time of the first retrieved chunk's metadata for this file
                        # This assumes all chunks from the same file version have the same original file mod time
                        stored_mod_time = existing_data['metadatas'][0].get('modified_at_timestamp')
                        if stored_mod_time == current_file_metadata['modified_at_timestamp']:
                            logger.info(f"File {file_name} has not been modified since last load (Timestamp: {stored_mod_time}). Skipping.")
                            needs_update = False
                        else:
                            logger.info(f"File {file_name} has been modified (Current: {current_file_metadata['modified_at_timestamp']}, Stored: {stored_mod_time}). Reprocessing.")
                    else:
                        logger.info(f"No existing chunks found for {file_name}. Will process as new.")
                except Exception as e:
                    logger.warning(f"Could not reliably check existing chunks for {file_name} due to: {e}. Will process.")


            if needs_update or force_reprocess_all_files:
                if not needs_update and force_reprocess_all_files: # Log if forced
                    logger.info(f"Force reprocessing enabled for {file_name}.")
                
                # If updating, delete old chunks for this file first
                if needs_update: # Also true if force_reprocess_all_files caused needs_update to remain true
                    try:
                        ids_to_delete = collection_to_load.get(where={"file_name": file_name}, include=[])['ids']
                        if ids_to_delete:
                            logger.info(f"Deleting {len(ids_to_delete)} old chunks for updated file: {file_name}")
                            collection_to_load.delete(ids=ids_to_delete)
                        else:
                             logger.info(f"No old chunks found to delete for {file_name} (might be new or already cleared).")
                    except Exception as e:
                        logger.error(f"Error deleting old chunks for {file_name}: {e}. Continuing with adding new chunks.")


                full_text = ""
                if file_suffix == ".pdf":
                    full_text = extract_text_from_pdf_with_ocr(file_path, force_ocr=force_ocr_all_pdfs)
                elif file_suffix == ".docx":
                    full_text = extract_text_from_docx(file_path)

                if not full_text or not full_text.strip():
                    logger.warning(f"No text extracted from {file_name}. Skipping.")
                    continue
                
                cleaned_text = clean_text(full_text)
                chunks = split_text_into_chunks(cleaned_text, file_name)
                
                if not chunks:
                    logger.warning(f"No chunks generated for {file_name} after splitting. Skipping.")
                    continue
                
                batch_documents, batch_metadatas, batch_ids = [], [], []
                for i, chunk_text in enumerate(chunks):
                    # Create a more robust unique ID, e.g., using file name and chunk index + timestamp
                    chunk_id = f"{file_name}_chunk_{i}_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
                    
                    chunk_meta = {
                        # Standard metadata for retrieval and filtering
                        "file_name": file_name, # Critical for identifying source file
                        "chunk_number": i,
                        "total_chunks_in_doc": len(chunks),
                        "chunk_char_length": len(chunk_text),
                        "chunk_token_count": count_tokens(chunk_text),
                        "chunk_preview": chunk_text[:150].strip().replace("\n", " ") + "...",
                        # Include all file-level metadata for potential future use or detailed inspection
                        **current_file_metadata 
                    }
                    batch_documents.append(chunk_text)
                    batch_metadatas.append(chunk_meta)
                    batch_ids.append(chunk_id)
                
                if batch_documents:
                    try:
                        logger.info(f"Adding {len(batch_documents)} chunks for {file_name} to ChromaDB.")
                        collection_to_load.add(documents=batch_documents, metadatas=batch_metadatas, ids=batch_ids)
                        logger.info(f"Successfully added/updated chunks for {file_name}.")
                        new_chunks_added_this_run += len(batch_documents)
                    except Exception as e:
                        logger.error(f"Error adding batch to ChromaDB for {file_name}: {e}")
            processed_files_count +=1
        else: # Is a directory or other non-file item
            logger.debug(f"Skipping item (not a file): {item.name}")


    logger.info("--- Document Loading Summary ---")
    logger.info(f"Total files checked/processed in this run: {processed_files_count}")
    logger.info(f"New chunks added/updated in collection in this run: {new_chunks_added_this_run}")
    try:
        total_chunks_in_collection = collection_to_load.count()
        logger.info(f"Total chunks now in '{COLLECTION_NAME}': {total_chunks_in_collection}")
    except Exception as e:
        logger.error(f"Could not get total count from collection '{COLLECTION_NAME}': {e}")

# ==============================================================================
# SECCIÓN 8: EJECUCIÓN PRINCIPAL
# ==============================================================================
# Este es el punto de entrada del script cuando se ejecuta directamente.
# Verifica la configuración, inicializa los componentes (ChromaDB, función de
# embedding) y llama a la función principal de procesamiento de documentos.
# Incluye lógica para crear un archivo de demostración si el directorio de
# entrada está vacío.

if __name__ == "__main__":
    logger.info("Starting HR Document Loader script with OCR capabilities.")
    
    if not OPENAI_API_KEY:
        logger.error("CRITICAL: OPENAI_API_KEY environment variable not set. This is required for embeddings. Exiting.")
        exit(1)
    
    if not HR_DOCS_DIR.is_dir():
        logger.info(f"HR Docs input directory not found at {HR_DOCS_DIR}. Creating it now.")
        HR_DOCS_DIR.mkdir(parents=True, exist_ok=True)
        logger.info(f"Please place your .pdf and .docx documents in {HR_DOCS_DIR} and re-run the script.")
        # exit(0) # Optionally exit if you want user to add files first

    # Example: Create a dummy docx if HR_DOCS_DIR is empty, for testing.
    if not any(HR_DOCS_DIR.iterdir()): # Check if directory is empty
        logger.info(f"{HR_DOCS_DIR} is empty. Creating a dummy .docx file for demonstration.")
        try:
            dummy_doc = Document()
            dummy_doc.add_heading("Sample Company Policy Document", 0)
            dummy_doc.add_paragraph(
                "This document outlines the sample policies for employees at OurCompany. "
                "It includes information on work hours, vacation, and benefits."
            )
            dummy_doc.add_heading("Work Hours", level=1)
            dummy_doc.add_paragraph("Standard work hours are 9 AM to 5 PM, Monday to Friday.")
            dummy_doc.add_heading("Vacation Policy", level=1)
            dummy_doc.add_paragraph(
                "Employees are entitled to 20 days of paid vacation per year, subject to approval. "
                "Requests should be submitted at least two weeks in advance through the HR portal."
            )
            dummy_doc.save(HR_DOCS_DIR / "dummy_policy_document.docx")
            logger.info(f"Created dummy file: {HR_DOCS_DIR / 'dummy_policy_document.docx'}")
        except Exception as e:
            logger.error(f"Could not create dummy .docx file: {e}")


    # You can add an argument parser here to control force_ocr_all_pdfs from command line
    # For example: `python loader.py --force-ocr --force-reprocess`
    FORCE_OCR_ALL_PDFS = False # Set to True to OCR all PDFs regardless of digital text extraction success
    FORCE_REPROCESS_ALL = False # Set to True to reprocess all files even if not modified

    try:
        openai_embedding_function = get_embedding_function()
        chroma_client_instance = initialize_chroma_client()
        hr_collection_instance = get_or_create_collection(chroma_client_instance, openai_embedding_function)
        
        process_and_load_documents(
            hr_collection_instance,
            force_ocr_all_pdfs=FORCE_OCR_ALL_PDFS,
            force_reprocess_all_files=FORCE_REPROCESS_ALL
        )
        
        logger.info("HR Document Loader script finished successfully.")
    
    except ValueError as ve: # For API key or config errors
        logger.error(f"Configuration error: {ve}")
    except Exception as e:
        logger.error(f"An unexpected error occurred during loader execution: {e}", exc_info=True)